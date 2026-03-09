"""embeddings — Shared vector-search library for ai-search, ai-chat, and ai-organize.

Consolidates database access, embedding generation, text extraction, chunking,
indexing, and retrieval into a single importable module.  Each consumer imports
only what it needs.

Not meant to be run standalone — no ``if __name__ == '__main__'`` block.
"""

from __future__ import annotations

import json
import mimetypes
import os
import sqlite3
import struct
import sys
import time
from pathlib import Path

import sqlite_vec

# ── Import Ollama client ──────────────────────────────────────────────────────
from ollama import embed as _ollama_embed, CHAT_MODEL, EMBED_MODEL  # noqa: F401

# ── Import config for pipeline settings ───────────────────────────────────────
from config import get as _cfg

# ── Configuration ─────────────────────────────────────────────────────────────

XDG_DATA_HOME = os.environ.get("XDG_DATA_HOME", os.path.expanduser("~/.local/share"))
APP_DIR = Path(XDG_DATA_HOME) / "ai-search"
DB_PATH = APP_DIR / "vectors.db"
FEEDBACK_DB_PATH = APP_DIR / "feedback.db"

EMBED_DIM = 1024
CHUNK_SIZE = 4000
OVERLAP = 200

# ── Hybrid Retrieval Constants ────────────────────────────────────────────────
RRF_K = int(_cfg("retrieval", "rrf_k", 60))
BM25_ENABLED = _cfg("retrieval", "bm25_enabled", True)
UTILITY_WEIGHT = float(_cfg("retrieval", "utility_weight", 0.15))

# Pure-text extensions that mimetypes might miss.
# Consumers can extend this set (ai-organize adds .xml, .ini, etc.).
TEXT_EXTS = {
    ".nix", ".md", ".toml", ".yml", ".yaml", ".sh", ".bash", ".zsh",
    ".json", ".txt", ".py", ".rs", ".go", ".js", ".ts", ".jsx", ".tsx",
    ".css", ".html", ".csv",
}

# Binary formats we know how to extract text from
BINARY_EXTS = {".pdf", ".docx", ".xlsx"}

# Directories to skip when walking trees
IGNORE_DIRS = {".git", "node_modules", "vendor", "__pycache__", ".venv", "dist", "build"}


# ── Database Layer ────────────────────────────────────────────────────────────

def init_db(db_path: Path | None = None) -> sqlite3.Connection:
    """Create (or open) the vector database with schema + sqlite-vec loaded.

    This is the *write-mode* opener used by indexing operations.  It creates
    the app directory, enables WAL, and ensures both tables exist.
    """
    db = db_path or DB_PATH
    db.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(db)
    conn.enable_load_extension(True)
    sqlite_vec.load(conn)
    conn.enable_load_extension(False)
    conn.execute("PRAGMA journal_mode=WAL;")

    conn.execute(f"""
        CREATE VIRTUAL TABLE IF NOT EXISTS file_embeddings USING vec0(
            embedding float[{EMBED_DIM}]
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS file_metadata (
            rowid INTEGER PRIMARY KEY,
            filepath TEXT NOT NULL,
            mtime REAL NOT NULL,
            snippet TEXT NOT NULL,
            chunk_text TEXT NOT NULL DEFAULT ''
        )
    """)
    # Migrate older databases that lack the chunk_text column
    try:
        conn.execute("SELECT chunk_text FROM file_metadata LIMIT 0")
    except sqlite3.OperationalError:
        conn.execute("ALTER TABLE file_metadata ADD COLUMN chunk_text TEXT NOT NULL DEFAULT ''")

    # ── FTS5 full-text index for hybrid BM25 + vector search ─────────────
    conn.execute("""
        CREATE VIRTUAL TABLE IF NOT EXISTS file_chunks_fts USING fts5(
            chunk_text,
            content='file_metadata',
            content_rowid='rowid'
        )
    """)

    # Triggers to keep FTS5 in sync with file_metadata inserts/deletes.
    # "IF NOT EXISTS" isn't supported for triggers, so use try/except.
    try:
        conn.execute("""
            CREATE TRIGGER file_metadata_ai AFTER INSERT ON file_metadata BEGIN
                INSERT INTO file_chunks_fts(rowid, chunk_text)
                VALUES (new.rowid, new.chunk_text);
            END
        """)
    except sqlite3.OperationalError:
        pass  # trigger already exists

    try:
        conn.execute("""
            CREATE TRIGGER file_metadata_ad AFTER DELETE ON file_metadata BEGIN
                INSERT INTO file_chunks_fts(file_chunks_fts, rowid, chunk_text)
                VALUES ('delete', old.rowid, old.chunk_text);
            END
        """)
    except sqlite3.OperationalError:
        pass  # trigger already exists

    try:
        conn.execute("""
            CREATE TRIGGER file_metadata_au AFTER UPDATE ON file_metadata BEGIN
                INSERT INTO file_chunks_fts(file_chunks_fts, rowid, chunk_text)
                VALUES ('delete', old.rowid, old.chunk_text);
                INSERT INTO file_chunks_fts(rowid, chunk_text)
                VALUES (new.rowid, new.chunk_text);
            END
        """)
    except sqlite3.OperationalError:
        pass  # trigger already exists

    # Rebuild the FTS index on first migration (populates from existing data).
    # This is a no-op if the FTS table already has content.
    try:
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM file_chunks_fts")
        fts_count = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM file_metadata")
        meta_count = cur.fetchone()[0]
        if meta_count > 0 and fts_count == 0:
            conn.execute("INSERT INTO file_chunks_fts(file_chunks_fts) VALUES('rebuild')")
            print("FTS5 index rebuilt from existing data.", file=sys.stderr)
    except sqlite3.OperationalError:
        pass  # vec0 table might confuse COUNT in some edge cases

    return conn


def open_db(db_path: Path | None = None) -> sqlite3.Connection:
    """Open an *existing* database in read mode (no schema creation).

    Raises SystemExit with a JSON error message if the DB doesn't exist.
    """
    db = db_path or DB_PATH
    if not db.exists():
        print(json.dumps({
            "error": (
                f"No search database found at {db}. "
                "Run: ai-search --index <directory>"
            )
        }))
        sys.exit(1)
    conn = sqlite3.connect(db)
    conn.enable_load_extension(True)
    sqlite_vec.load(conn)
    conn.enable_load_extension(False)
    return conn


def clear_db(db_path: Path | None = None) -> None:
    """Delete the database and its WAL/SHM companions."""
    db = db_path or DB_PATH
    if db.exists():
        db.unlink()
        for ext in (".db-wal", ".db-shm"):
            sidecar = db.with_suffix(ext)
            if sidecar.exists():
                sidecar.unlink()
        print("Database cleared.", file=sys.stderr)
    else:
        print("Database not found. Nothing to clear.", file=sys.stderr)


# ── Embedding Helpers ─────────────────────────────────────────────────────────

def get_embedding(
    text: str,
    *,
    model: str | None = None,
    timeout: int = 10,
    keep_alive: int | None = None,
) -> list[float]:
    """Fetch the embedding vector from Ollama for *text*.

    Delegates to :func:`ollama.embed`. See its docstring for parameter details.
    """
    return _ollama_embed(text, model=model, timeout=timeout, keep_alive=keep_alive)


def vec_to_bytes(vec: list[float]) -> bytes:
    """Pack a float vector into the raw-bytes format sqlite-vec expects."""
    return struct.pack(f"<{len(vec)}f", *vec)


def bytes_to_vec(raw: bytes) -> list[float]:
    """Unpack raw bytes from sqlite-vec into a float list."""
    n = len(raw) // 4
    return list(struct.unpack(f"<{n}f", raw))


# ── Text Extraction & Chunking ───────────────────────────────────────────────

def is_text_file(filepath: Path, *, extra_text_exts: set[str] | None = None) -> bool:
    """Check if a file is a supported text or binary format.

    Parameters
    ----------
    extra_text_exts : set[str], optional
        Additional extensions to treat as plain text (e.g. ai-organize passes
        {".xml", ".ini", ".cfg", ".conf", ".env", ".r", ".sql", ".jl"}).
    """
    ext = filepath.suffix.lower()
    all_text = TEXT_EXTS | (extra_text_exts or set())
    if ext in all_text or ext in BINARY_EXTS:
        return True
    mime, _ = mimetypes.guess_type(str(filepath))
    return mime is not None and mime.startswith("text/")


def extract_text(path: Path) -> str | None:
    """Extract plain text from a file, dispatching on extension for binary formats."""
    ext = path.suffix.lower()

    # Plain-text formats — read directly
    if ext not in BINARY_EXTS:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            return None

    # PDF — extract page-by-page, prefixed with page numbers for context
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"[Page {i}]\n{text.strip()}")
            return "\n\n".join(pages) if pages else None
        except Exception:
            return None

    # DOCX — extract paragraphs and table cells
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts) if parts else None
        except Exception:
            return None

    # XLSX — extract each sheet with its name and tab-separated rows
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"[Sheet: {sheet_name}]")
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join("" if c is None else str(c) for c in row)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts) if parts else None
        except Exception:
            return None

    return None


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = OVERLAP) -> list[str]:
    """Split text into overlapping chunks."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += (chunk_size - overlap)
    return chunks


# ── Indexing ──────────────────────────────────────────────────────────────────

def index_directory(
    conn: sqlite3.Connection,
    directory: str,
    *,
    embed_timeout: int = 10,
    extra_text_exts: set[str] | None = None,
    extra_ignore_dirs: set[str] | None = None,
) -> tuple[int, int]:
    """Walk *directory*, embed the text files, and store them.

    Returns (files_indexed, chunks_indexed).
    """
    all_ignore = IGNORE_DIRS | (extra_ignore_dirs or set())
    root_dir = Path(directory).resolve()
    if not root_dir.is_dir():
        print(f"Error: {root_dir} is not a valid directory.", file=sys.stderr)
        sys.exit(1)

    print(f"Indexing {root_dir}...", file=sys.stderr)
    indexed_files = 0
    indexed_chunks = 0

    cur = conn.cursor()
    cur.execute("SELECT filepath, mtime FROM file_metadata")
    existing_files = {row[0]: row[1] for row in cur.fetchall()}

    for path in root_dir.rglob("*"):
        if any(part in all_ignore for part in path.parts):
            continue

        if not path.is_file() or not is_text_file(path, extra_text_exts=extra_text_exts):
            continue

        try:
            mtime = path.stat().st_mtime
        except OSError:
            continue

        filepath_str = str(path)

        # Skip if file hasn't been modified since we last indexed it
        if filepath_str in existing_files and existing_files[filepath_str] >= mtime:
            continue

        content = extract_text(path)
        if content is None:
            continue

        # If it was modified, delete the old chunks first
        if filepath_str in existing_files:
            cur.execute("SELECT rowid FROM file_metadata WHERE filepath = ?", (filepath_str,))
            rowids = [r[0] for r in cur.fetchall()]
            if rowids:
                placeholders = ",".join("?" * len(rowids))
                cur.execute(f"DELETE FROM file_embeddings WHERE rowid IN ({placeholders})", rowids)
                cur.execute(f"DELETE FROM file_metadata WHERE rowid IN ({placeholders})", rowids)

        chunks = chunk_text(content)
        for chunk in chunks:
            snippet = " ".join(chunk[:200].split())
            if not snippet:
                continue

            embedding = get_embedding(chunk, timeout=embed_timeout)
            if not embedding:
                continue

            vec_bytes = vec_to_bytes(embedding)

            cur.execute("""
                INSERT INTO file_metadata (filepath, mtime, snippet, chunk_text)
                VALUES (?, ?, ?, ?)
            """, (filepath_str, mtime, snippet, chunk))

            new_id = cur.lastrowid
            cur.execute("""
                INSERT INTO file_embeddings (rowid, embedding)
                VALUES (?, ?)
            """, (new_id, vec_bytes))
            indexed_chunks += 1

        indexed_files += 1
        if indexed_files % 10 == 0:
            conn.commit()
            print(f"... Indexed {indexed_files} modified files so far.", file=sys.stderr)

    conn.commit()
    print(f"Indexing complete. Added {indexed_files} new/modified files ({indexed_chunks} chunks).",
          file=sys.stderr)
    return indexed_files, indexed_chunks


def check_dir_indexed(conn: sqlite3.Connection, directory: str) -> bool:
    """Return True if *directory* has at least one indexed file."""
    abs_dir = str(Path(directory).resolve())
    prefix = abs_dir if abs_dir.endswith("/") else abs_dir + "/"
    cur = conn.cursor()
    try:
        cur.execute("SELECT 1 FROM file_metadata WHERE filepath LIKE ? LIMIT 1", (prefix + "%",))
        return cur.fetchone() is not None
    except sqlite3.OperationalError:
        return False


# ── Retrieval ─────────────────────────────────────────────────────────────────

def _vector_search(
    conn: sqlite3.Connection,
    query_vec_bytes: bytes,
    *,
    top_k: int = 15,
    max_distance: float = 0.8,
    scope: str | None = None,
) -> list[dict]:
    """Raw vector search — returns ranked results with rowid for RRF fusion."""
    cur = conn.cursor()
    fetch_limit = top_k * 4 if scope else top_k
    scope_prefix = (scope.rstrip("/") + "/") if scope else None

    try:
        cur.execute("""
            SELECT
                m.rowid,
                m.filepath,
                m.snippet,
                COALESCE(m.chunk_text, m.snippet) AS chunk_text,
                vec_distance_cosine(e.embedding, ?) AS distance
            FROM file_embeddings e
            JOIN file_metadata m ON e.rowid = m.rowid
            ORDER BY distance ASC
            LIMIT ?
        """, (query_vec_bytes, fetch_limit))
        rows = cur.fetchall()
    except sqlite3.OperationalError as e:
        print(json.dumps({"error": f"Vector search failed: {e}"}))
        return []

    results = []
    for rowid, filepath, snippet, chunk_text, dist in rows:
        if dist > max_distance:
            continue
        if scope_prefix and not filepath.startswith(scope_prefix):
            continue
        score = round(max(0.0, 1.0 - dist), 4)
        results.append({
            "rowid": rowid,
            "filepath": filepath,
            "snippet": snippet,
            "chunk_text": chunk_text,
            "score": score,
            "distance": dist,
        })
        if len(results) >= top_k:
            break

    return results


def bm25_search(
    conn: sqlite3.Connection,
    query: str,
    *,
    top_k: int = 20,
    scope: str | None = None,
) -> list[dict]:
    """Keyword search over chunk_text using SQLite FTS5 BM25 ranking.

    Returns results with rowid for RRF fusion.  Gracefully returns [] if the
    FTS5 table doesn't exist or the query has no keyword matches.
    """
    scope_prefix = (scope.rstrip("/") + "/") if scope else None

    # FTS5 MATCH requires at least one indexable term.  Sanitise the query
    # to avoid syntax errors from stray quotes or operators.
    fts_query = _sanitize_fts_query(query)
    if not fts_query:
        return []

    cur = conn.cursor()
    try:
        if scope_prefix:
            cur.execute("""
                SELECT m.rowid, m.filepath, m.snippet,
                       COALESCE(m.chunk_text, m.snippet) AS chunk_text,
                       rank
                FROM file_chunks_fts fts
                JOIN file_metadata m ON fts.rowid = m.rowid
                WHERE file_chunks_fts MATCH ?
                  AND m.filepath LIKE ?
                ORDER BY rank
                LIMIT ?
            """, (fts_query, scope_prefix + "%", top_k))
        else:
            cur.execute("""
                SELECT m.rowid, m.filepath, m.snippet,
                       COALESCE(m.chunk_text, m.snippet) AS chunk_text,
                       rank
                FROM file_chunks_fts fts
                JOIN file_metadata m ON fts.rowid = m.rowid
                WHERE file_chunks_fts MATCH ?
                ORDER BY rank
                LIMIT ?
            """, (fts_query, top_k))

        return [
            {
                "rowid": r[0],
                "filepath": r[1],
                "snippet": r[2],
                "chunk_text": r[3],
                "bm25_rank": i + 1,
                "score": 0.0,  # placeholder; RRF will assign the real score
            }
            for i, r in enumerate(cur.fetchall())
        ]
    except sqlite3.OperationalError:
        # FTS5 table might not exist yet (read-only open_db path)
        return []


def _sanitize_fts_query(query: str) -> str:
    """Turn a user query into a safe FTS5 MATCH expression.

    Strips FTS5 operators and wraps each word in quotes so that literal
    terms are matched without triggering syntax errors.
    """
    import re as _re
    # Remove FTS5 special chars: AND, OR, NOT, quotes, parens, colons, *, ^
    words = _re.findall(r"[a-zA-Z0-9_./-]+", query)
    if not words:
        return ""
    # Join as implicit OR by quoting each term
    return " OR ".join(f'"{w}"' for w in words)


def reciprocal_rank_fusion(
    vec_results: list[dict],
    bm25_results: list[dict],
    utility_scores: dict[int, float] | None = None,
    *,
    k: int = RRF_K,
    utility_weight: float = UTILITY_WEIGHT,
) -> list[dict]:
    """Fuse vector and BM25 ranked lists using Reciprocal Rank Fusion.

    Each result dict must have a ``rowid`` key (used for dedup across the two
    lists).  When *utility_scores* is provided, each chunk's accumulated
    utility EMA is blended in as a third ranking signal.
    """
    scores: dict[int, float] = {}
    lookup: dict[int, dict] = {}

    for i, r in enumerate(vec_results):
        rid = r["rowid"]
        scores[rid] = scores.get(rid, 0.0) + 1.0 / (k + i + 1)
        lookup[rid] = r

    for i, r in enumerate(bm25_results):
        rid = r["rowid"]
        scores[rid] = scores.get(rid, 0.0) + 1.0 / (k + i + 1)
        lookup.setdefault(rid, r)

    # Blend in learned utility scores from the feedback store
    if utility_scores:
        for rid, ema in utility_scores.items():
            if rid in scores:
                scores[rid] += utility_weight * ema

    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [lookup[rid] for rid, _ in ranked]


def search(
    conn: sqlite3.Connection,
    query: str,
    *,
    top_k: int = 5,
    max_distance: float = 0.8,
    embed_timeout: int = 10,
) -> list[dict]:
    """Embed *query* and return the top-K closest chunks as dicts.

    Each dict has keys: filepath, snippet, distance.
    """
    embedding = get_embedding(query, timeout=embed_timeout)
    if not embedding:
        sys.exit(1)

    vec_bytes = vec_to_bytes(embedding)
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT
                m.filepath,
                m.snippet,
                vec_distance_cosine(e.embedding, ?) as distance
            FROM file_embeddings e
            JOIN file_metadata m ON e.rowid = m.rowid
            ORDER BY distance ASC
            LIMIT ?
        """, (vec_bytes, top_k))
        results = cur.fetchall()
    except sqlite3.OperationalError as e:
        print(f"Error: Could not query database ({e})", file=sys.stderr)
        sys.exit(1)

    return [
        {"filepath": fp, "snippet": snippet, "distance": round(dist, 4)}
        for fp, snippet, dist in results
        if dist <= max_distance
    ]


def retrieve_with_chunks(
    conn: sqlite3.Connection,
    query: str,
    *,
    top_k: int = 5,
    max_distance: float = 0.8,
    scope: str | None = None,
    embed_timeout: int = 30,
) -> list[dict]:
    """Retrieve top-K chunks with full chunk_text for RAG context injection.

    Uses hybrid BM25 + vector search with Reciprocal Rank Fusion when
    BM25 is enabled (default).  Falls back to pure vector search when
    BM25 is disabled or the FTS5 table is unavailable.

    Each dict has keys: rowid, filepath, snippet, chunk_text, score.
    Used by ai-chat for grounded generation.
    """
    embedding = get_embedding(query, timeout=embed_timeout)
    if not embedding:
        return []

    vec_bytes = vec_to_bytes(embedding)
    widen_k = top_k * 3

    # ── Vector search (always runs) ──────────────────────────────────────
    vec_results = _vector_search(
        conn, vec_bytes, top_k=widen_k,
        max_distance=max_distance, scope=scope,
    )

    # ── BM25 search (hybrid path, when enabled) ─────────────────────────
    if BM25_ENABLED:
        bm25_results = bm25_search(conn, query, top_k=widen_k, scope=scope)
    else:
        bm25_results = []

    # ── Load utility scores from feedback store (if available) ───────────
    utility = _load_utility_scores()

    if bm25_results or utility:
        fused = reciprocal_rank_fusion(vec_results, bm25_results, utility)
    else:
        fused = vec_results

    return fused[:top_k]


def _load_utility_scores() -> dict[int, float] | None:
    """Load chunk utility scores from the feedback DB with temporal decay.

    Returns None if the feedback DB doesn't exist or has no data.
    This is a read-only operation — safe to call from any context.
    """
    if not FEEDBACK_DB_PATH.exists():
        return None

    decay_halflife = float(_cfg("learning", "decay_halflife_days", 30))
    now = time.time()

    try:
        fconn = sqlite3.connect(FEEDBACK_DB_PATH)
        cur = fconn.cursor()
        cur.execute("SELECT chunk_rowid, utility_ema, last_used FROM chunk_utility")
        rows = cur.fetchall()
        fconn.close()
    except sqlite3.OperationalError:
        return None

    if not rows:
        return None

    scores: dict[int, float] = {}
    for rid, ema, last_used in rows:
        age_days = (now - last_used) / 86400
        decay = 0.5 ** (age_days / decay_halflife) if decay_halflife > 0 else 1.0
        # Decay toward neutral (0.5), not toward zero
        decayed = 0.5 + (ema - 0.5) * decay
        scores[rid] = decayed

    return scores


# ── Cache Helpers (for ai-organize) ──────────────────────────────────────────

def load_cached_embeddings(directory: str, db_path: Path | None = None) -> dict[str, list[float]] | None:
    """Load averaged embeddings from the vector DB for *directory*.

    Returns a dict mapping relative file paths to float vectors, or None if
    the DB doesn't exist or has no data for this directory.  Multiple chunks
    per file are averaged into a single representative vector.
    """
    db = db_path or DB_PATH
    if not db.exists():
        print("  No ai-search DB found — will embed from scratch.", file=sys.stderr)
        return None

    abs_dir = str(Path(directory).resolve())
    prefix = abs_dir.rstrip("/") + "/"

    try:
        conn = sqlite3.connect(db)
        conn.enable_load_extension(True)
        sqlite_vec.load(conn)
        conn.enable_load_extension(False)
        cur = conn.cursor()
        cur.execute(
            "SELECT m.filepath, e.embedding "
            "FROM file_metadata m "
            "JOIN file_embeddings e ON m.rowid = e.rowid "
            "WHERE m.filepath LIKE ?",
            (prefix + "%",),
        )
        rows = cur.fetchall()
        conn.close()
    except Exception as e:
        print(f"  Could not read ai-search DB for cache: {e}", file=sys.stderr)
        return None

    if not rows:
        print("  ai-search DB exists but has no embeddings for this directory.",
              file=sys.stderr)
        return None

    # Average multiple chunks per file into one vector
    file_vecs: dict[str, list[float]] = {}
    file_cnts: dict[str, int] = {}
    vec_len: int | None = None

    for filepath, vec_bytes in rows:
        if vec_bytes is None:
            continue
        rel = filepath.removeprefix(prefix)
        if vec_len is None:
            vec_len = len(vec_bytes) // 4
        vec = bytes_to_vec(vec_bytes)
        if rel in file_vecs:
            file_vecs[rel] = [a + b for a, b in zip(file_vecs[rel], vec)]
            file_cnts[rel] += 1
        else:
            file_vecs[rel] = vec
            file_cnts[rel] = 1

    if not file_vecs:
        return None

    for rel in file_vecs:
        cnt = file_cnts[rel]
        if cnt > 1:
            file_vecs[rel] = [v / cnt for v in file_vecs[rel]]

    return file_vecs


def save_embeddings(
    embeddings: list[tuple[str, list[float]]],
    db_path: Path | None = None,
) -> None:
    """Write freshly computed embeddings back to the vector DB.

    Each entry is (absolute_filepath, embedding_vector).  Existing rows for
    the same filepath are replaced so repeated runs stay idempotent.

    Unlike the chunked indexing path, this stores one row per file with the
    filename as the snippet — suitable for ai-organize's metadata embeddings.
    """
    if not embeddings:
        return

    db = db_path or DB_PATH
    try:
        db.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(db)
        conn.enable_load_extension(True)
        sqlite_vec.load(conn)
        conn.enable_load_extension(False)
        conn.execute("PRAGMA journal_mode=WAL;")

        # Ensure tables exist (mirrors the canonical schema)
        conn.execute(f"""
            CREATE VIRTUAL TABLE IF NOT EXISTS file_embeddings USING vec0(
                embedding float[{EMBED_DIM}]
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS file_metadata (
                rowid INTEGER PRIMARY KEY,
                filepath TEXT NOT NULL,
                mtime REAL NOT NULL,
                snippet TEXT NOT NULL
            )
        """)

        cur = conn.cursor()
        written = 0

        for filepath, vec in embeddings:
            # Skip zero-vectors
            if all(v == 0.0 for v in vec[:8]):
                continue

            try:
                mtime = Path(filepath).stat().st_mtime
            except OSError:
                mtime = 0.0

            # Delete existing rows for this filepath
            cur.execute("SELECT rowid FROM file_metadata WHERE filepath = ?", (filepath,))
            old_rowids = [r[0] for r in cur.fetchall()]
            if old_rowids:
                placeholders = ",".join("?" * len(old_rowids))
                cur.execute(f"DELETE FROM file_embeddings WHERE rowid IN ({placeholders})", old_rowids)
                cur.execute(f"DELETE FROM file_metadata WHERE rowid IN ({placeholders})", old_rowids)

            snippet = Path(filepath).name
            cur.execute(
                "INSERT INTO file_metadata (filepath, mtime, snippet) VALUES (?, ?, ?)",
                (filepath, mtime, snippet),
            )
            new_id = cur.lastrowid

            vec_bytes = vec_to_bytes(vec)
            cur.execute(
                "INSERT INTO file_embeddings (rowid, embedding) VALUES (?, ?)",
                (new_id, vec_bytes),
            )
            written += 1

            if written % 100 == 0:
                conn.commit()

        conn.commit()
        conn.close()
        print(f"  Saved {written} new embeddings to ai-search DB.", file=sys.stderr)

    except Exception as e:
        print(f"  Could not write embeddings to ai-search DB: {e}", file=sys.stderr)


# ── Status ────────────────────────────────────────────────────────────────────

def get_status(conn: sqlite3.Connection) -> dict:
    """Return database statistics as a dict."""
    cur = conn.cursor()
    try:
        cur.execute("SELECT COUNT(DISTINCT filepath), COUNT(*) FROM file_metadata")
        files, chunks = cur.fetchone()
        size_mb = DB_PATH.stat().st_size / (1024 * 1024)
        return {
            "db_path": str(DB_PATH),
            "size_mb": round(size_mb, 2),
            "files_indexed": files or 0,
            "total_chunks": chunks or 0,
        }
    except sqlite3.OperationalError:
        return {
            "db_path": str(DB_PATH),
            "size_mb": 0,
            "files_indexed": 0,
            "total_chunks": 0,
        }

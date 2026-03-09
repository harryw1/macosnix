#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "sqlite-vec",
#     "requests",
#     "pypdf",
#     "python-docx",
#     "openpyxl",
# ]
# ///
import argparse
import json
import mimetypes
import os
import sqlite3
import struct
import sys
import time
from pathlib import Path

import requests
import sqlite_vec

# Determine XDG Data Home
XDG_DATA_HOME = os.environ.get("XDG_DATA_HOME", os.path.expanduser("~/.local/share"))
APP_DIR = Path(XDG_DATA_HOME) / "ai-search"
DB_PATH = APP_DIR / "vectors.db"

OLLAMA_EMBED_URL = "http://localhost:11434/api/embeddings"
MODEL = os.environ.get("OLLAMA_MODEL_EMBED", "qwen3-embedding:0.6b")

# Maximum characters per chunk to avoid hitting token limits
CHUNK_SIZE = 4000
OVERLAP = 200

# Additional pure-text extensions that mimetypes might miss
TEXT_EXTS = {".nix", ".md", ".toml", ".yml", ".yaml", ".sh", ".bash", ".zsh", ".json", ".txt", ".py", ".rs", ".go", ".js", ".ts", ".jsx", ".tsx", ".css", ".html", ".csv"}
# Binary formats we know how to extract text from
BINARY_EXTS = {".pdf", ".docx", ".xlsx"}
IGNORE_DIRS = {".git", "node_modules", "vendor", "__pycache__", ".venv", "dist", "build"}


def get_embedding(text: str) -> list[float]:
    """Fetch the embedding vector from Ollama for the given text."""
    try:
        resp = requests.post(
            OLLAMA_EMBED_URL,
            json={"model": MODEL, "prompt": text},
            timeout=10,
        )
        resp.raise_for_status()
        return resp.json().get("embedding", [])
    except requests.exceptions.RequestException as e:
        print(f"Error fetching embedding from Ollama: {e}", file=sys.stderr)
        sys.exit(1)


def is_text_file(filepath: Path) -> bool:
    """Check if a file is a supported text or binary format."""
    ext = filepath.suffix.lower()
    if ext in TEXT_EXTS or ext in BINARY_EXTS:
        return True
    mime, _ = mimetypes.guess_type(str(filepath))
    return mime is not None and mime.startswith("text/")


def extract_text(path: Path) -> str | None:
    """Extract plain text from a file, dispatching on extension for binary formats."""
    ext = path.suffix.lower()

    # Plain-text formats — read directly
    if ext in TEXT_EXTS:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            return None

    # PDF — extract page-by-page, prefixed with page numbers for context
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"[Page {i}]\n{text.strip()}")
            return "\n\n".join(pages) if pages else None
        except Exception:
            return None

    # DOCX — extract paragraphs and table cells
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts) if parts else None
        except Exception:
            return None

    # XLSX — extract each sheet with its name and tab-separated rows
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"[Sheet: {sheet_name}]")
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join("" if c is None else str(c) for c in row)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts) if parts else None
        except Exception:
            return None

    return None


def chunk_text(text: str, chunk_size=CHUNK_SIZE, overlap=OVERLAP) -> list[str]:
    """Split text into overlapping chunks."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += (chunk_size - overlap)
    return chunks


def init_db() -> sqlite3.Connection:
    """Initialize the SQLite database with sqlite-vec extension."""
    APP_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.enable_load_extension(True)
    sqlite_vec.load(conn)
    conn.enable_load_extension(False)
    
    # Enable Write-Ahead Logging for better concurrent performance
    conn.execute("PRAGMA journal_mode=WAL;")

    # We use a virtual table provided by sqlite-vec.
    # qwen3-embedding:0.6b produces 1024-dimensional vectors.
    conn.execute("""
        CREATE VIRTUAL TABLE IF NOT EXISTS file_embeddings USING vec0(
            embedding float[1024]
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS file_metadata (
            rowid INTEGER PRIMARY KEY,
            filepath TEXT NOT NULL,
            mtime REAL NOT NULL,
            snippet TEXT NOT NULL
        )
    """)
    return conn


def index_directory(conn: sqlite3.Connection, directory: str):
    """Walk a directory, embed the text files, and store them."""
    root_dir = Path(directory).resolve()
    if not root_dir.is_dir():
        print(f"Error: {root_dir} is not a valid directory.", file=sys.stderr)
        sys.exit(1)

    print(f"Indexing {root_dir}...", file=sys.stderr)
    indexed_files = 0
    indexed_chunks = 0
    
    # Prepare checking existing files to avoid re-indexing unchanged files
    cur = conn.cursor()
    cur.execute("SELECT filepath, mtime FROM file_metadata")
    existing_files = {row[0]: row[1] for row in cur.fetchall()}

    for path in root_dir.rglob("*"):
        # Ignore common bad directories
        if any(part in IGNORE_DIRS for part in path.parts):
            continue
            
        if not path.is_file() or not is_text_file(path):
            continue

        try:
            mtime = path.stat().st_mtime
        except OSError:
            continue

        filepath_str = str(path)
        
        # Skip if file hasn't been modified since we last indexed it
        if filepath_str in existing_files and existing_files[filepath_str] >= mtime:
            continue

        content = extract_text(path)
        if content is None:
            continue

        # If it was modified, delete the old chunks first
        if filepath_str in existing_files:
            cur.execute("SELECT rowid FROM file_metadata WHERE filepath = ?", (filepath_str,))
            rowids = [r[0] for r in cur.fetchall()]
            if rowids:
                placeholders = ",".join("?" * len(rowids))
                cur.execute(f"DELETE FROM file_embeddings WHERE rowid IN ({placeholders})", rowids)
                cur.execute(f"DELETE FROM file_metadata WHERE rowid IN ({placeholders})", rowids)
        
        chunks = chunk_text(content)
        for chunk in chunks:
            # We save a cleaned 200 char snippet for context
            snippet = " ".join(chunk[:200].split())
            if not snippet:
                continue
                
            embedding = get_embedding(chunk)
            if not embedding:
               continue
               
            # sqlite-vec expects raw bytes for float vectors
            vec_bytes = struct.pack(f"<{len(embedding)}f", *embedding)
            
            # Insert metadata and get rowid
            cur.execute("""
                INSERT INTO file_metadata (filepath, mtime, snippet)
                VALUES (?, ?, ?)
            """, (filepath_str, mtime, snippet))
            
            # Insert vector exactly at the same rowid
            new_id = cur.lastrowid
            cur.execute("""
                INSERT INTO file_embeddings (rowid, embedding)
                VALUES (?, ?)
            """, (new_id, vec_bytes))
            indexed_chunks += 1
            
        indexed_files += 1
        # Periodically commit and log
        if indexed_files % 10 == 0:
            conn.commit()
            print(f"... Indexed {indexed_files} modified files so far.", file=sys.stderr)

    conn.commit()
    print(f"Indexing complete. Added {indexed_files} new/modified files ({indexed_chunks} chunks).", file=sys.stderr)


def search(conn: sqlite3.Connection, query: str):
    """Embed the search query and find the top K closest vectors."""
    embedding = get_embedding(query)
    if not embedding:
        sys.exit(1)
        
    vec_bytes = struct.pack(f"<{len(embedding)}f", *embedding)
    
    # We want top 5 results
    cur = conn.cursor()
    try:
        # Distance natively returned by sqlite-vec. Smaller distance = closer.
        cur.execute("""
            SELECT 
                m.filepath, 
                m.snippet,
                vec_distance_cosine(e.embedding, ?) as distance
            FROM file_embeddings e
            JOIN file_metadata m ON e.rowid = m.rowid
            ORDER BY distance ASC
            LIMIT 5
        """, (vec_bytes,))
        results = cur.fetchall()
    except sqlite3.OperationalError as e:
         print(f"[]\nError: Could not query database. Have you indexed anything yet? ({e})", file=sys.stderr)
         sys.exit(1)

    output = []
    # sqlite-vec returns cosine distance. 0 is identical, 2 is opposite.
    for r in results:
        filepath, snippet, dist = r
        # Optional: Skip results that are completely unrelated (distance > 0.8 is a good threshold for this model)
        if dist > 0.8:
            continue
        output.append({
            "filepath": filepath,
            "snippet": snippet,
            "distance": round(dist, 4)
        })
        
    print(json.dumps(output))


def get_status(conn: sqlite3.Connection):
    """Print database statistics."""
    cur = conn.cursor()
    try:
        cur.execute("SELECT COUNT(DISTINCT filepath), COUNT(*) FROM file_metadata")
        files, chunks = cur.fetchone()
        size_mb = DB_PATH.stat().st_size / (1024 * 1024)
        print(json.dumps({
            "db_path": str(DB_PATH),
            "size_mb": round(size_mb, 2),
            "files_indexed": files or 0,
            "total_chunks": chunks or 0,
        }))
    except sqlite3.OperationalError:
         print(json.dumps({
            "db_path": str(DB_PATH),
            "size_mb": 0,
            "files_indexed": 0,
            "total_chunks": 0,
        }))


def check_dir(conn: sqlite3.Connection, directory: str):
    """Exit 0 if the directory has indexed files, else exit 1."""
    abs_dir = str(Path(directory).resolve())
    prefix = abs_dir if abs_dir.endswith("/") else abs_dir + "/"
    cur = conn.cursor()
    try:
        cur.execute("SELECT 1 FROM file_metadata WHERE filepath LIKE ? LIMIT 1", (prefix + "%",))
        if cur.fetchone():
            sys.exit(0)
        else:
            sys.exit(1)
    except sqlite3.OperationalError:
        sys.exit(1)


def clear_db():
    """Delete the database completely."""
    if DB_PATH.exists():
        DB_PATH.unlink()
        wal = DB_PATH.with_suffix(".db-wal")
        if wal.exists(): wal.unlink()
        shm = DB_PATH.with_suffix(".db-shm")
        if shm.exists(): shm.unlink()
        print("Database cleared.", file=sys.stderr)
    else:
        print("Database not found. Nothing to clear.", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="Local Semantic Search Python Backend")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--index", metavar="DIR", help="Directory to recursively index")
    group.add_argument("--search", metavar="QUERY", help="Semantic search query")
    group.add_argument("--status", action="store_true", help="Print database statistics as JSON")
    group.add_argument("--clear", action="store_true", help="Delete the vector database")
    group.add_argument("--check-dir", metavar="DIR", help="Check if directory is indexed (exits 0 if true, 1 if false)")
    
    args = parser.parse_args()

    if args.clear:
        clear_db()
        sys.exit(0)

    # All other commands require the database connection
    conn = init_db()

    if args.status:
        get_status(conn)
    elif args.index:
        index_directory(conn, args.index)
    elif args.search:
        search(conn, args.search)
    elif args.check_dir:
        check_dir(conn, args.check_dir)

    conn.close()


if __name__ == "__main__":
    main()

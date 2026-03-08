#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "markdown>=3.6",
#   "python-docx>=1.1",
#   "weasyprint>=60",
#   "beautifulsoup4>=4.12",
#   "pyyaml>=6",
# ]
# ///
"""
mdconvert – Markdown → docx / html / pdf

Styled to match Harry's professional document templates:
  • report : Deep Navy #203354 · Blue-Grey #4A6572 · Gold accent #F9AA33
  • notes  : Clean black/grey, Times New Roman, simple

Usage:
  mdconvert -f docx  report.md            →  report.docx
  mdconvert -f pdf   report.md            →  report.pdf
  mdconvert -f html  report.md            →  report.html
  mdconvert -f pdf  -t notes meeting.md  →  meeting.pdf
  mdconvert report.md report.docx         # explicit output path

Front-matter (YAML between --- delimiters) supports:
  title, subtitle, author, email, organization, date, template
"""

import argparse
import re
import sys
from datetime import date
from io import BytesIO
from pathlib import Path

import yaml
import markdown as md_lib
from bs4 import BeautifulSoup, NavigableString, Tag


# ── Colour palettes ────────────────────────────────────────────────────────────

PALETTES: dict[str, dict[str, str]] = {
    "report": {
        "primary":       "#203354",
        "secondary":     "#4A6572",
        "accent":        "#F9AA33",
        "text":          "#333333",
        "code_bg":       "#F5F5F5",
        "code_border":   "#E0E0E0",
        "table_head_bg": "#203354",
        "table_head_fg": "#FFFFFF",
        "table_alt_row": "#F2F4F6",
        "table_border":  "#D0D7DE",
        "rule_color":    "#4A6572",
        "quote_color":   "#4A6572",
        "link_color":    "#203354",
    },
    "notes": {
        "primary":       "#1A1A1A",
        "secondary":     "#444444",
        "accent":        "#1565C0",
        "text":          "#1A1A1A",
        "code_bg":       "#F5F5F5",
        "code_border":   "#D0D0D0",
        "table_head_bg": "#444444",
        "table_head_fg": "#FFFFFF",
        "table_alt_row": "#F7F7F7",
        "table_border":  "#D0D0D0",
        "rule_color":    "#AAAAAA",
        "quote_color":   "#666666",
        "link_color":    "#1565C0",
    },
}

# ── Front-matter ───────────────────────────────────────────────────────────────

_YAML_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)

DEFAULT_META: dict[str, str] = {
    "author":       "Harrison Weiss",
    "email":        "harrisonrweiss1@gmail.com",
    "organization": "",
    "date":         date.today().strftime("%B %-d, %Y"),
}


def parse_front_matter(text: str) -> tuple[dict, str]:
    m = _YAML_RE.match(text)
    if not m:
        return {}, text
    meta = yaml.safe_load(m.group(1)) or {}
    return meta, text[m.end():]


# ── Markdown → HTML ────────────────────────────────────────────────────────────

_MD_EXTENSIONS = [
    "tables",
    "fenced_code",
    "footnotes",
    "attr_list",
    "def_list",
    "nl2br",
    "sane_lists",
    "toc",
    "smarty",
]


def md_to_html(body: str) -> str:
    # Pre-process GFM strikethrough: ~~text~~ → <del>text</del>
    # The standard markdown library doesn't support this natively; we handle it
    # with a regex that only matches within a single line (won't touch code
    # fences, which are processed by the fenced_code extension afterward).
    body = re.sub(r"~~([^\n~]+)~~", r"<del>\1</del>", body)
    return md_lib.markdown(body, extensions=_MD_EXTENSIONS)


# ── Colour helpers ─────────────────────────────────────────────────────────────

def _hex(color: str) -> str:
    """Strip leading # from a hex colour string."""
    return color.lstrip("#")


def _rgb(color: str):
    """Convert #RRGGBB to docx RGBColor."""
    from docx.shared import RGBColor
    h = _hex(color)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ── docx XML helpers ───────────────────────────────────────────────────────────

def _para_bottom_border(para, color: str, sz: int = 8) -> None:
    """Add a bottom border line under a paragraph (replicates LaTeX \\titlerule)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _hex(color))
    pBdr.append(bot)
    pPr.append(pBdr)


def _shade_para(para, fill: str) -> None:
    """Apply background shading to an entire paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(fill))
    pPr.append(shd)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(fill))
    tcPr.append(shd)


def _cell_borders(cell, color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), _hex(color))
        borders.append(b)
    tcPr.append(borders)


def _add_hyperlink(para, url: str, text: str, palette: dict) -> None:
    """Insert a properly linked hyperlink run into a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), _hex(palette["link_color"]))
    rPr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    hl.append(r)
    para._p.append(hl)


# ── Inline walker ──────────────────────────────────────────────────────────────

def _walk_inline(
    para,
    node,
    palette: dict,
    bold: bool = False,
    italic: bool = False,
    strike: bool = False,
    mono: bool = False,
) -> None:
    """Recursively add inline content (text, bold, italic, code, links…) to a paragraph."""
    from docx.shared import Pt

    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = para.add_run(text)
            run.bold = bold or None
            run.italic = italic or None
            run.font.strike = strike or None
            if mono:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        return

    tag = node.name if isinstance(node, Tag) else None

    if tag == "br":
        para.add_run("\n")
        return

    if tag == "a":
        href = node.get("href", "")
        link_text = node.get_text()
        if href and link_text:
            _add_hyperlink(para, href, link_text, palette)
        elif link_text:
            run = para.add_run(link_text)
            run.font.color.rgb = _rgb(palette["link_color"])
            run.underline = True
        return

    new_bold   = bold   or tag in ("strong", "b")
    new_italic = italic or tag in ("em", "i")
    new_strike = strike or tag in ("del", "s", "strike")
    new_mono   = mono   or tag == "code"

    for child in node.children:
        _walk_inline(para, child, palette, new_bold, new_italic, new_strike, new_mono)


# ── List walker ────────────────────────────────────────────────────────────────

def _walk_list(doc, list_node, palette: dict, template: str, level: int = 0, ordered: bool = False) -> None:
    from docx.shared import Pt, Inches

    style = "List Number" if ordered else "List Bullet"

    for child in list_node.children:
        if not isinstance(child, Tag) or child.name != "li":
            continue

        inline_parts: list = []
        nested: list[Tag] = []
        for item in child.children:
            if isinstance(item, Tag) and item.name in ("ul", "ol"):
                nested.append(item)
            else:
                inline_parts.append(item)

        para = doc.add_paragraph(style=style)
        para.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
        para.paragraph_format.space_after  = Pt(2)
        para.paragraph_format.space_before = Pt(0)
        for part in inline_parts:
            # Skip whitespace-only text nodes — BeautifulSoup leaves a trailing
            # "\n" after a nested <ul>/<ol> inside an <li>, and _walk_inline
            # would turn it into a <w:br/>, producing a spurious blank line.
            if isinstance(part, NavigableString) and not str(part).strip():
                continue
            _walk_inline(para, part, palette)

        for nl in nested:
            _walk_list(doc, nl, palette, template, level + 1, ordered=(nl.name == "ol"))


# ── Table builder ──────────────────────────────────────────────────────────────

_PAGE_WIDTH_INCHES = 6.5  # 8.5 in page − 1 in left − 1 in right margin


def _build_table(doc, table_node: Tag, palette: dict) -> None:
    from docx.shared import Pt, Inches

    # Collect header + body row nodes (Tag objects, not plain text, so we can
    # walk inline formatting — bold/italic/code — inside cells)
    thead = table_node.find("thead")
    tbody = table_node.find("tbody")

    header_rows: list[list[Tag]] = []
    body_rows:   list[list[Tag]] = []

    if thead:
        for tr in thead.find_all("tr"):
            header_rows.append(tr.find_all(["th", "td"]))
    if tbody:
        for tr in tbody.find_all("tr"):
            body_rows.append(tr.find_all(["th", "td"]))

    # Fallback: table has no explicit thead/tbody (first <tr> becomes header)
    if not header_rows and not body_rows:
        all_trs = table_node.find_all("tr")
        if not all_trs:
            return
        header_rows = [all_trs[0].find_all(["th", "td"])]
        body_rows   = [tr.find_all(["th", "td"]) for tr in all_trs[1:]]

    all_rows = header_rows + body_rows
    n_head   = len(header_rows)
    if not all_rows:
        return

    n_cols = max(len(r) for r in all_rows)
    n_rows = len(all_rows)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Distribute columns evenly across the full text width (fixes wide-table
    # overflow: each column gets an equal share of the 6.5 in text area)
    col_width = Inches(_PAGE_WIDTH_INCHES / n_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    for ri, cell_nodes in enumerate(all_rows):
        is_header = ri < n_head
        body_idx  = ri - n_head

        row = table.rows[ri]
        for ci in range(n_cols):
            cell = row.cells[ci]
            _cell_borders(cell, palette["table_border"])

            if is_header:
                _shade_cell(cell, palette["table_head_bg"])
                para = cell.paragraphs[0]
                if ci < len(cell_nodes):
                    # Walk inline so code/emphasis inside headers is preserved,
                    # then force bold + white colour on every resulting run.
                    for child in cell_nodes[ci].children:
                        _walk_inline(para, child, palette)
                    for run in para.runs:
                        run.bold           = True
                        run.font.size      = Pt(10)
                        run.font.color.rgb = _rgb(palette["table_head_fg"])
            else:
                if body_idx % 2 == 1:
                    _shade_cell(cell, palette["table_alt_row"])
                para = cell.paragraphs[0]
                if ci < len(cell_nodes):
                    # Full inline walk: preserves bold, italic, code, links
                    for child in cell_nodes[ci].children:
                        _walk_inline(para, child, palette)
                    for run in para.runs:
                        if run.font.size is None:
                            run.font.size = Pt(10)

    # Spacing after table
    space_para = doc.add_paragraph()
    space_para.paragraph_format.space_after  = Pt(8)
    space_para.paragraph_format.space_before = Pt(0)


# ── Block walker ───────────────────────────────────────────────────────────────

def _walk_block(doc, parent, palette: dict, template: str) -> None:
    from docx.shared import Pt, Inches

    for node in parent.children:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                para = doc.add_paragraph(text)
                para.paragraph_format.space_after = Pt(6)
            continue

        if not isinstance(node, Tag):
            continue

        tag = node.name

        # ── Headings ────────────────────────────────────────────────────────
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            para  = doc.add_paragraph()

            spacings = {1: (18, 6), 2: (14, 4), 3: (10, 4), 4: (8, 2), 5: (6, 2), 6: (4, 2)}
            pre, post = spacings.get(level, (6, 4))
            para.paragraph_format.space_before = Pt(pre)
            para.paragraph_format.space_after  = Pt(post)

            run = para.add_run(node.get_text(strip=True))

            if level == 1:
                run.font.size      = Pt(20)
                run.bold           = True
                run.font.color.rgb = _rgb(palette["primary"])
                _para_bottom_border(para, palette["secondary"], sz=8)
            elif level == 2:
                run.font.size      = Pt(16)
                run.bold           = True
                run.font.color.rgb = _rgb(palette["secondary"])
            elif level == 3:
                run.font.size      = Pt(13)
                run.bold           = True
                run.font.color.rgb = _rgb(palette["secondary"])
            else:
                run.font.size      = Pt(11)
                run.bold           = True
                run.font.color.rgb = _rgb(palette["text"])

        # ── Paragraph ───────────────────────────────────────────────────────
        elif tag == "p":
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            for child in node.children:
                _walk_inline(para, child, palette)

        # ── Lists ────────────────────────────────────────────────────────────
        elif tag in ("ul", "ol"):
            _walk_list(doc, node, palette, template, level=0, ordered=(tag == "ol"))

        # ── Blockquote ───────────────────────────────────────────────────────
        elif tag == "blockquote":
            for child in node.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if not text:
                        continue
                    para = doc.add_paragraph()
                else:
                    para = doc.add_paragraph()
                    for gc in child.children:
                        _walk_inline(para, gc, palette)

                para.paragraph_format.left_indent  = Inches(0.4)
                para.paragraph_format.space_after   = Pt(4)
                # Colour existing runs
                for run in para.runs:
                    run.font.color.rgb = _rgb(palette["quote_color"])
                    run.italic = True

        # ── Fenced code block ────────────────────────────────────────────────
        elif tag == "pre":
            code_node = node.find("code")
            text = (code_node or node).get_text()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
            para.paragraph_format.left_indent  = Inches(0.2)
            run = para.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            _shade_para(para, palette["code_bg"])

        # ── Horizontal rule ──────────────────────────────────────────────────
        elif tag == "hr":
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after  = Pt(8)
            _para_bottom_border(para, palette["rule_color"], sz=6)

        # ── Table ────────────────────────────────────────────────────────────
        elif tag == "table":
            _build_table(doc, node, palette)

        # ── Containers ───────────────────────────────────────────────────────
        elif tag in ("div", "section", "article", "body", "main"):
            _walk_block(doc, node, palette, template)

        # ── Anything else ────────────────────────────────────────────────────
        else:
            text = node.get_text(strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.paragraph_format.space_after = Pt(6)


# ── DOCX builder ───────────────────────────────────────────────────────────────

def build_docx(meta: dict, html: str, palette: dict, template: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Page margins (1 in on all sides)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Title / header block ─────────────────────────────────────────────────
    if meta.get("title"):
        if template == "report":
            # Decorative top rule
            rule = doc.add_paragraph()
            rule.paragraph_format.space_before = Pt(0)
            rule.paragraph_format.space_after  = Pt(10)
            _para_bottom_border(rule, palette["secondary"], sz=18)

            # Title
            tp = doc.add_paragraph()
            tr = tp.add_run(meta["title"])
            tr.font.size      = Pt(26)
            tr.bold           = True
            tr.font.color.rgb = _rgb(palette["primary"])
            tp.paragraph_format.space_before = Pt(4)
            tp.paragraph_format.space_after  = Pt(4)

            # Subtitle
            if meta.get("subtitle"):
                sp = doc.add_paragraph()
                sr = sp.add_run(meta["subtitle"])
                sr.font.size      = Pt(15)
                sr.font.color.rgb = _rgb(palette["secondary"])
                sp.paragraph_format.space_after = Pt(8)

            # Date
            dp = doc.add_paragraph()
            dr = dp.add_run(str(meta.get("date", DEFAULT_META["date"])))
            dr.font.size      = Pt(11)
            dr.font.color.rgb = _rgb(palette["text"])
            dp.paragraph_format.space_after = Pt(4)

            # Author / org
            parts = []
            author = meta.get("author", DEFAULT_META["author"])
            if author:
                parts.append(f"Prepared by: {author}")
            org = meta.get("organization", DEFAULT_META["organization"])
            if org:
                parts.append(org)
            if parts:
                ip = doc.add_paragraph()
                ir = ip.add_run("  ·  ".join(parts))
                ir.font.size      = Pt(10)
                ir.font.color.rgb = _rgb(palette["secondary"])
                ip.paragraph_format.space_after = Pt(18)

            doc.add_page_break()

        else:  # notes
            tp = doc.add_paragraph()
            tr = tp.add_run(meta["title"])
            tr.font.size      = Pt(20)
            tr.bold           = True
            _para_bottom_border(tp, palette["secondary"], sz=6)
            tp.paragraph_format.space_after = Pt(4)

            # Date · Author line
            parts = [str(meta.get("date", DEFAULT_META["date"]))]
            author = meta.get("author", DEFAULT_META["author"])
            if author:
                parts.append(author)
            ip = doc.add_paragraph()
            ir = ip.add_run("  ·  ".join(parts))
            ir.font.size      = Pt(10)
            ir.font.color.rgb = _rgb(palette["secondary"])
            ip.paragraph_format.space_after = Pt(12)

    # ── Body ─────────────────────────────────────────────────────────────────
    soup = BeautifulSoup(html, "html.parser")
    _walk_block(doc, soup, palette, template)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTML / CSS builder ─────────────────────────────────────────────────────────

def _css(palette: dict, template: str) -> str:
    p = palette
    if template == "report":
        body_font    = "Palatino, Georgia, 'Times New Roman', serif"
        heading_font = "'Arial', 'Helvetica Neue', Helvetica, sans-serif"
        mono_font    = "'JetBrains Mono', 'Courier New', Courier, monospace"
    else:
        body_font    = "'Times New Roman', Times, serif"
        heading_font = "'Arial', 'Helvetica Neue', Helvetica, sans-serif"
        mono_font    = "'Courier New', Courier, monospace"

    return f"""
/* ── Reset ─────────────────────────────────────────────── */
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}

/* ── Base ──────────────────────────────────────────────── */
body {{
  font-family: {body_font};
  font-size: 11pt;
  line-height: 1.65;
  color: {p['text']};
  max-width: 820px;
  margin: 0 auto;
  padding: 2.5rem 3rem;
}}

/* ── Title block ────────────────────────────────────────── */
.title-block {{
  margin-bottom: 2.2rem;
  padding-bottom: 1.2rem;
  border-bottom: 2px solid {p['secondary']};
}}
.title-block .doc-title {{
  font-family: {heading_font};
  font-size: 26pt;
  font-weight: 700;
  color: {p['primary']};
  line-height: 1.15;
  margin-bottom: 0.25em;
}}
.title-block .doc-subtitle {{
  font-family: {heading_font};
  font-size: 15pt;
  color: {p['secondary']};
  margin-bottom: 0.5em;
}}
.title-block .doc-meta {{
  font-size: 10pt;
  color: {p['secondary']};
  margin-top: 0.35em;
}}

/* Notes variant */
.title-block.notes .doc-title {{
  font-size: 20pt;
  color: {p['primary']};
  border-bottom: 1px solid {p['secondary']};
  padding-bottom: 0.3em;
  margin-bottom: 0.3em;
}}

/* ── Headings ───────────────────────────────────────────── */
h1, h2, h3, h4, h5, h6 {{
  font-family: {heading_font};
  line-height: 1.3;
  margin-top: 1.8em;
  margin-bottom: 0.55em;
}}
h1 {{
  font-size: 19pt;
  color: {p['primary']};
  border-bottom: 2px solid {p['secondary']};
  padding-bottom: 0.2em;
}}
h2 {{ font-size: 15pt; color: {p['secondary']}; }}
h3 {{ font-size: 12pt; color: {p['secondary']}; }}
h4, h5, h6 {{ font-size: 11pt; color: {p['text']}; }}

/* ── Body text ──────────────────────────────────────────── */
p {{ margin-bottom: 0.85em; }}

/* ── Links ──────────────────────────────────────────────── */
a {{ color: {p['link_color']}; text-decoration: underline; }}
a:hover {{ opacity: 0.8; }}

/* ── Tables ─────────────────────────────────────────────── */
table {{
  border-collapse: collapse;
  width: 100%;
  margin: 1.1em 0;
  font-size: 10pt;
}}
thead tr {{
  background: {p['table_head_bg']};
  color: {p['table_head_fg']};
}}
th, td {{
  padding: 8px 13px;
  text-align: left;
  border: 1px solid {p['table_border']};
  vertical-align: top;
}}
th {{
  font-family: {heading_font};
  font-weight: 700;
  font-size: 9.5pt;
  letter-spacing: 0.02em;
}}
tbody tr:nth-child(even) {{ background: {p['table_alt_row']}; }}

/* ── Code ───────────────────────────────────────────────── */
code {{
  font-family: {mono_font};
  font-size: 9pt;
  background: {p['code_bg']};
  border: 1px solid {p['code_border']};
  padding: 0.1em 0.35em;
  border-radius: 3px;
}}
pre {{
  background: {p['code_bg']};
  border: 1px solid {p['code_border']};
  border-left: 3px solid {p['secondary']};
  padding: 1em 1.2em;
  overflow-x: auto;
  margin: 1em 0;
  border-radius: 0 4px 4px 0;
  font-size: 9pt;
}}
pre code {{
  background: none;
  border: none;
  padding: 0;
  border-radius: 0;
  line-height: 1.55;
}}

/* ── Blockquote ─────────────────────────────────────────── */
blockquote {{
  border-left: 3px solid {p['quote_color']};
  margin: 1em 0;
  padding: 0.5em 1.1em;
  color: {p['quote_color']};
  font-style: italic;
}}

/* ── Lists ──────────────────────────────────────────────── */
ul, ol {{ margin: 0.5em 0 0.9em 1.6em; padding: 0; }}
li {{ margin-bottom: 0.25em; }}
li > ul, li > ol {{ margin-top: 0.25em; }}

/* ── HR ─────────────────────────────────────────────────── */
hr {{
  border: none;
  border-top: 1px solid {p['rule_color']};
  margin: 1.6em 0;
}}

/* ── Print / PDF ────────────────────────────────────────── */
@media print {{
  body {{ padding: 0; max-width: none; font-size: 10.5pt; }}
  pre {{ white-space: pre-wrap; word-break: break-all; }}
  a {{ color: {p['link_color']}; }}
  h1, h2, h3 {{ page-break-after: avoid; }}
  table, pre, blockquote {{ page-break-inside: avoid; }}
  @page {{ margin: 1in; }}
}}
"""


def build_html(meta: dict, html_body: str, palette: dict, template: str) -> str:
    title_str = meta.get("title", "")
    css_str   = _css(palette, template)

    # Build title block
    tb = ""
    if title_str:
        cls = "title-block notes" if template == "notes" else "title-block"
        sub = (
            f'<div class="doc-subtitle">{meta["subtitle"]}</div>'
            if meta.get("subtitle") else ""
        )
        date_str   = str(meta.get("date", DEFAULT_META["date"]))
        author_str = meta.get("author", DEFAULT_META["author"])
        org_str    = meta.get("organization", DEFAULT_META["organization"])

        if template == "report":
            meta_parts = [date_str]
            if author_str:
                meta_parts.append(f"Prepared by: {author_str}")
            if org_str:
                meta_parts.append(org_str)
            meta_html = "  ·  ".join(meta_parts)
            tb = f"""
<div class="{cls}">
  <div class="doc-title">{title_str}</div>
  {sub}
  <div class="doc-meta">{meta_html}</div>
</div>"""
        else:  # notes
            meta_parts = [date_str]
            if author_str:
                meta_parts.append(author_str)
            meta_html = "  ·  ".join(meta_parts)
            tb = f"""
<div class="{cls}">
  <div class="doc-title">{title_str}</div>
  <div class="doc-meta">{meta_html}</div>
</div>"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title_str}</title>
  <style>{css_str}</style>
</head>
<body>
{tb}
<div class="content">
{html_body}
</div>
</body>
</html>"""


# ── PDF builder ────────────────────────────────────────────────────────────────

def build_pdf(meta: dict, html_body: str, palette: dict, template: str) -> bytes:
    from weasyprint import HTML
    full_html = build_html(meta, html_body, palette, template)
    return HTML(string=full_html).write_pdf()


# ── CLI ────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        prog="mdconvert",
        description="Markdown → docx / html / pdf with professional styling.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  mdconvert -f docx  report.md              →  report.docx
  mdconvert -f pdf   report.md              →  report.pdf
  mdconvert -f html  report.md              →  report.html
  mdconvert -f pdf  -t notes  meeting.md    →  meeting.pdf  (meeting-notes style)
  mdconvert report.md report.docx           # explicit output path

Front-matter (YAML between --- delimiters):
  title, subtitle, author, email, organization, date, template
""",
    )
    parser.add_argument("input",            help="Input Markdown file")
    parser.add_argument("output", nargs="?", help="Output file (format inferred from extension)")
    parser.add_argument("-f", "--format",   choices=["docx", "pdf", "html"],
                        help="Output format (inferred from output extension if omitted)")
    parser.add_argument("-t", "--template", choices=["report", "notes"], default=None,
                        help="Document style: 'report' (default) or 'notes'")
    parser.add_argument("-o", "--out",      help="Output file path (alternative to positional arg)")

    args = parser.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        print(f"mdconvert: error: '{in_path}' not found", file=sys.stderr)
        sys.exit(1)

    raw          = in_path.read_text(encoding="utf-8")
    meta, body   = parse_front_matter(raw)

    # Fill in defaults
    for k, v in DEFAULT_META.items():
        meta.setdefault(k, v)

    # Resolve output path and format
    out_path_str = args.out or args.output
    if out_path_str:
        out_path = Path(out_path_str)
        fmt      = args.format or out_path.suffix.lstrip(".")
    elif args.format:
        fmt      = args.format
        out_path = in_path.with_suffix(f".{fmt}")
    else:
        parser.error("specify --format/-f or provide an output file path")
        return  # unreachable but appeases type checkers

    if fmt not in ("docx", "pdf", "html"):
        print(f"mdconvert: error: unknown format '{fmt}'", file=sys.stderr)
        sys.exit(1)

    # Template: CLI flag > front-matter key > default
    template = args.template or meta.get("template", "report")
    palette  = PALETTES[template]

    # Convert
    html_body = md_to_html(body)

    if fmt == "docx":
        out_path.write_bytes(build_docx(meta, html_body, palette, template))
    elif fmt == "html":
        out_path.write_text(build_html(meta, html_body, palette, template), encoding="utf-8")
    elif fmt == "pdf":
        out_path.write_bytes(build_pdf(meta, html_body, palette, template))

    print(f"✓  {out_path}")


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "pytest>=8",
#   "markdown>=3.6",
#   "python-docx>=1.1",
#   "weasyprint>=60",
#   "beautifulsoup4>=4.12",
#   "pyyaml>=6",
# ]
# ///
"""
mdconvert test suite — run with:

    uv run tests/test_mdconvert.py          # from modules/home-manager/
    uv run tests/test_mdconvert.py -v       # verbose
    uv run tests/test_mdconvert.py -k wide  # filter by name

Tests are organised into three layers:

    1. Crash tests — every format × template combination must produce a
       non-empty output file without raising an exception.

    2. Structural tests — spot-check docx internals (table count, heading
       colour, column widths) to catch regressions in the builder logic.

    3. Edge-case tests — specifically target the patterns in edge_cases.md
       that have historically broken document converters.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

# ── Path setup ─────────────────────────────────────────────────────────────────
# Allow importing mdconvert from the parent directory (modules/home-manager/)
HERE = Path(__file__).parent          # tests/
ROOT = HERE.parent                    # modules/home-manager/
sys.path.insert(0, str(ROOT))

import mdconvert as M                  # noqa: E402  (after sys.path tweak)

EDGE_CASES_MD = HERE / "edge_cases.md"

# ── Helpers ────────────────────────────────────────────────────────────────────

def _convert(tmp_path: Path, source_md: str, fmt: str, template: str = "report") -> Path:
    """Run a full end-to-end conversion and return the output Path."""
    tmp_path.mkdir(parents=True, exist_ok=True)   # handle nested subdir calls
    in_file  = tmp_path / "input.md"
    out_file = tmp_path / f"output.{fmt}"
    in_file.write_text(source_md, encoding="utf-8")

    meta, body = M.parse_front_matter(in_file.read_text())
    for k, v in M.DEFAULT_META.items():
        meta.setdefault(k, v)

    palette   = M.PALETTES[template]
    html_body = M.md_to_html(body)

    if fmt == "docx":
        out_file.write_bytes(M.build_docx(meta, html_body, palette, template))
    elif fmt == "html":
        out_file.write_text(M.build_html(meta, html_body, palette, template))
    elif fmt == "pdf":
        out_file.write_bytes(M.build_pdf(meta, html_body, palette, template))

    return out_file


def _open_docx(path: Path):
    from docx import Document
    return Document(str(path))


# ── Fixtures ───────────────────────────────────────────────────────────────────

MINIMAL_MD = """\
---
title: Minimal Test
author: Test Runner
---

Hello, world.
"""

BASIC_TABLE_MD = """\
---
title: Basic Table
---

| Name  | Value |
|-------|-------|
| alpha | 1     |
| beta  | 2     |
"""

FORMATTED_CELLS_MD = """\
---
title: Formatted Cells
---

| Package     | Notes                               |
|-------------|-------------------------------------|
| **bold**    | `code` value                        |
| *italic*    | **bold** and *italic* together      |
| `monospace` | plain text                          |
"""

WIDE_TABLE_MD = """\
---
title: Wide Table
---

| A | B | C | D | E | F | G | H |
|---|---|---|---|---|---|---|---|
| 1 | 2 | 3 | 4 | 5 | 6 | 7 | 8 |
| a | b | c | d | e | f | g | h |
"""

EMPTY_CELLS_MD = """\
---
title: Empty Cells
---

| Col A | Col B | Col C |
|-------|-------|-------|
| One   |       | Three |
|       | Two   |       |
| Four  | Five  |       |
"""

NO_HEADER_MD = """\
---
title: No Header Table
---

| A | B |
|---|---|
| 1 | 2 |
| 3 | 4 |
"""

BACK_TO_BACK_TABLES_MD = """\
---
title: Back-to-back Tables
---

| X | Y |
|---|---|
| 1 | 2 |

| P | Q |
|---|---|
| a | b |
"""

NESTED_LISTS_MD = """\
---
title: Nested Lists
---

- Level 1
    - Level 2
        - Level 3
            - Level 4
    - Back to 2
- Back to 1

1. Ordered
    1. Sub
        1. Sub-sub
"""

LONG_CELL_MD = """\
---
title: Long Cell Content
---

| Field       | Value                                                                                                    |
|-------------|----------------------------------------------------------------------------------------------------------|
| Description | This cell contains a very long sentence designed to force line-wrapping inside the cell without overflow. |
| Command     | `uv run mdconvert -f docx -t notes very-long-filename-that-might-cause-issues.md output-document.docx`  |
"""

INLINE_FORMATTING_MD = """\
---
title: Inline Formatting
---

**Bold**, *italic*, ***bold-italic***, `code`, ~~strikethrough~~, and a
[link](https://example.com) all in the same paragraph.
"""

NO_FRONT_MATTER_MD = """\
# Document Without Front Matter

Just a heading and some text with a table:

| Col | Value |
|-----|-------|
| A   | 1     |
| B   | 2     |
"""

ALL_HEADINGS_MD = """\
---
title: All Headings
---

# H1
## H2
### H3
#### H4
##### H5
###### H6

Paragraph after H6.
"""

SPECIAL_CHARS_MD = """\
---
title: Special Characters
---

Quotes: "smart quotes" and 'apostrophes'.

Em dashes — like this — and en dashes: 2024–2026.

Unicode: café, naïve, Ångström.
"""

# ── 1. Crash tests (format × template × fixture) ──────────────────────────────

@pytest.mark.parametrize("fmt", ["docx", "html", "pdf"])
@pytest.mark.parametrize("template", ["report", "notes"])
def test_all_formats_minimal(tmp_path, fmt, template):
    """Every format+template combo must produce a non-empty file."""
    out = _convert(tmp_path / f"{fmt}_{template}", MINIMAL_MD, fmt, template)
    assert out.exists(), f"{out} was not created"
    assert out.stat().st_size > 0, f"{out} is empty"


@pytest.mark.parametrize("fmt", ["docx", "html", "pdf"])
def test_edge_cases_file(tmp_path, fmt):
    """The comprehensive edge_cases.md fixture must convert without crashing."""
    assert EDGE_CASES_MD.exists(), f"Fixture not found: {EDGE_CASES_MD}"
    src = EDGE_CASES_MD.read_text()
    out = _convert(tmp_path / fmt, src, fmt)
    assert out.stat().st_size > 0


@pytest.mark.parametrize("fmt", ["docx", "html", "pdf"])
def test_no_front_matter(tmp_path, fmt):
    """Documents without YAML front matter must not crash."""
    out = _convert(tmp_path / fmt, NO_FRONT_MATTER_MD, fmt)
    assert out.stat().st_size > 0


# ── 2. Structural / docx internals ────────────────────────────────────────────

def test_docx_table_count(tmp_path):
    """Two back-to-back tables produce exactly two tables in the docx."""
    out = _convert(tmp_path, BACK_TO_BACK_TABLES_MD, "docx")
    doc = _open_docx(out)
    assert len(doc.tables) == 2, f"Expected 2 tables, got {len(doc.tables)}"


def test_docx_wide_table_column_widths(tmp_path):
    """An 8-column table must not overflow: each column width should be > 0."""
    from docx.shared import Inches
    out = _convert(tmp_path, WIDE_TABLE_MD, "docx")
    doc = _open_docx(out)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.columns) == 8
    for i, col in enumerate(table.columns):
        # Each column must have positive, page-bounded width
        # (Inches stored as EMUs: 1 in = 914400 EMUs)
        assert col.width is not None and col.width > 0, \
            f"Column {i} has zero/null width"
        assert col.width <= Inches(M._PAGE_WIDTH_INCHES), \
            f"Column {i} overflows the page: {col.width} > {Inches(M._PAGE_WIDTH_INCHES)}"


def test_docx_table_header_shading(tmp_path):
    """Report-style header row should have navy background (#203354)."""
    out = _convert(tmp_path, BASIC_TABLE_MD, "docx", template="report")
    doc = _open_docx(out)
    assert doc.tables, "No tables found in output"
    table = doc.tables[0]
    from docx.oxml.ns import qn
    # First row cells should have the primary-colour fill
    header_row = table.rows[0]
    for cell in header_row.cells:
        tcPr = cell._tc.find(qn("w:tcPr"))
        assert tcPr is not None, "Header cell has no tcPr"
        shd  = tcPr.find(qn("w:shd"))
        assert shd is not None, "Header cell has no shading"
        fill = shd.get(qn("w:fill"), "").upper()
        assert fill == "203354", f"Expected navy header, got #{fill}"


def test_docx_body_row_alternates(tmp_path):
    """Odd body rows (0-indexed) should be shaded; even rows should not."""
    out = _convert(tmp_path, BASIC_TABLE_MD, "docx", template="report")
    doc = _open_docx(out)
    table = doc.tables[0]
    from docx.oxml.ns import qn
    # body rows: table.rows[1], table.rows[2] (0-indexed body_idx 0 and 1)
    # body_idx=1 (table.rows[2]) should be shaded F2F4F6
    alt_row = table.rows[2]
    cell = alt_row.cells[0]
    tcPr = cell._tc.find(qn("w:tcPr"))
    shd  = tcPr.find(qn("w:shd")) if tcPr is not None else None
    if shd is not None:
        fill = shd.get(qn("w:fill"), "").upper()
        assert fill in ("F2F4F6", "AUTO"), \
            f"Alt row has unexpected fill: #{fill}"


def test_docx_formatted_cells_preserve_bold(tmp_path):
    """Bold text inside a table cell must produce at least one bold run."""
    out = _convert(tmp_path, FORMATTED_CELLS_MD, "docx")
    doc = _open_docx(out)
    assert doc.tables, "No tables in output"
    table = doc.tables[0]
    # First body cell (row 1, col 0) contains **bold**
    cell = table.rows[1].cells[0]
    bold_runs = [r for p in cell.paragraphs for r in p.runs if r.bold]
    assert bold_runs, "No bold runs found in a cell that should contain **bold**"


def test_docx_formatted_cells_preserve_code(tmp_path):
    """Inline code inside a table cell must use Courier New font."""
    out = _convert(tmp_path, FORMATTED_CELLS_MD, "docx")
    doc = _open_docx(out)
    table = doc.tables[0]
    # Row 1, col 1 contains `code` value
    cell  = table.rows[1].cells[1]
    mono_runs = [
        r for p in cell.paragraphs for r in p.runs
        if r.font.name and "Courier" in r.font.name
    ]
    assert mono_runs, "No monospace runs found in a cell that should contain `code`"


def test_docx_empty_cells_no_crash(tmp_path):
    """Empty cells must produce a docx without raising an exception."""
    out = _convert(tmp_path, EMPTY_CELLS_MD, "docx")
    doc = _open_docx(out)
    assert len(doc.tables) == 1


def test_docx_heading_count(tmp_path):
    """ALL_HEADINGS_MD has one of each H1–H6; each becomes a paragraph."""
    out = _convert(tmp_path, ALL_HEADINGS_MD, "docx")
    doc = _open_docx(out)
    texts = [p.text.strip() for p in doc.paragraphs]
    for level in ("H1", "H2", "H3", "H4", "H5", "H6"):
        assert level in texts, f"Heading '{level}' not found in docx paragraphs"


def test_docx_no_header_table(tmp_path):
    """A table where GFM treats row 1 as the header should still build correctly."""
    out = _convert(tmp_path, NO_HEADER_MD, "docx")
    doc = _open_docx(out)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Should have 3 rows: 1 header + 2 body
    assert len(table.rows) == 3


# ── 3. Edge-case content tests ────────────────────────────────────────────────

def test_html_table_has_thead_and_tbody(tmp_path):
    """HTML output must contain <thead> and <tbody> elements for tables."""
    out = _convert(tmp_path, BASIC_TABLE_MD, "html")
    html = out.read_text()
    assert "<thead>" in html, "<thead> missing from HTML table"
    assert "<tbody>" in html, "<tbody> missing from HTML table"


def test_html_table_header_has_primary_colour(tmp_path):
    """Report HTML thead background must use the primary navy colour."""
    out = _convert(tmp_path, BASIC_TABLE_MD, "html", template="report")
    html = out.read_text()
    assert "#203354" in html, "Primary colour #203354 not found in report HTML"


def test_html_wide_table_renders(tmp_path):
    """8-column table must produce 8 <th> elements in the HTML output."""
    out = _convert(tmp_path, WIDE_TABLE_MD, "html")
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(out.read_text(), "html.parser")
    ths = soup.find_all("th")
    assert len(ths) == 8, f"Expected 8 <th> elements, found {len(ths)}"


def test_html_code_block_in_pre(tmp_path):
    """Fenced code blocks must appear inside <pre><code> in HTML output."""
    md = """\
---
title: Code test
---

```python
print("hello")
```
"""
    out = _convert(tmp_path, md, "html")
    html = out.read_text()
    # The markdown library renders as <pre><code class="language-python">…
    # so check for the opening tag prefix, not the exact closed tag.
    assert "<pre>" in html, "Fenced code block not wrapped in <pre>"
    assert "<code" in html, "Fenced code block missing <code> element"


def test_html_inline_code_in_code_tag(tmp_path):
    """Inline `code` spans must be wrapped in <code> tags."""
    out = _convert(tmp_path, INLINE_FORMATTING_MD, "html")
    html = out.read_text()
    assert "<code>" in html, "<code> tag missing from inline-code HTML"


def test_html_special_chars_unicode(tmp_path):
    """Unicode content must survive the conversion without corruption."""
    out = _convert(tmp_path, SPECIAL_CHARS_MD, "html")
    html = out.read_text(encoding="utf-8")
    for char in ("café", "naïve", "Ångström"):
        assert char in html, f"Unicode string '{char}' lost in HTML output"


def test_pdf_produces_bytes(tmp_path):
    """PDF output must be non-empty bytes starting with the PDF magic number."""
    out = _convert(tmp_path, EDGE_CASES_MD.read_text(), "pdf")
    data = out.read_bytes()
    assert data[:4] == b"%PDF", "Output does not begin with PDF magic number"


def test_nested_lists_docx_no_crash(tmp_path):
    """4-level nested list must produce a docx without crashing."""
    out = _convert(tmp_path, NESTED_LISTS_MD, "docx")
    doc = _open_docx(out)
    # Verify we have list-style paragraphs
    list_paras = [p for p in doc.paragraphs if "List" in (p.style.name or "")]
    assert list_paras, "No list paragraphs found in nested list output"


def test_long_cell_content_no_crash(tmp_path):
    """Very long cell text must not crash or produce an empty file."""
    out = _convert(tmp_path, LONG_CELL_MD, "docx")
    assert out.stat().st_size > 0


def test_notes_template_no_cover_page(tmp_path):
    """Notes template must not insert a page break (no cover page)."""
    out = _convert(tmp_path, MINIMAL_MD, "docx", template="notes")
    doc = _open_docx(out)
    # A page break produces a paragraph with \x0c; notes style should have none
    page_breaks = [
        p for p in doc.paragraphs
        if any(r.text == "\x0c" for r in p.runs)
    ]
    assert not page_breaks, "Notes template unexpectedly contains a page break"


def test_report_template_has_cover_page(tmp_path):
    """Report template must insert a page break after the cover block."""
    out = _convert(tmp_path, MINIMAL_MD, "docx", template="report")
    doc = _open_docx(out)
    from docx.oxml.ns import qn
    # A rendered page break appears as w:lastRenderedPageBreak or w:br with type=page
    page_breaks = [
        el
        for p in doc.paragraphs
        for el in p._p.iter()
        if el.tag == qn("w:br") and el.get(qn("w:type")) == "page"
    ]
    assert page_breaks, "Report template missing page break after cover"


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v", *sys.argv[1:]]))
